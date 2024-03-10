import fitz  # PyMuPDF
import json
from elasticsearch import Elasticsearch
import os
from docx import Document
from bert import BERT
import torch
import numpy as np
import uuid
import re

# Initialize Elasticsearch
ES_HOSTS = json.loads(os.environ.get('ES_HOSTS', '["https://ragllm-es01-1:9200", "https://ragllm-es02-1:9200", "https://ragllm-es03-1:9200"]'))
ES_CA_CERTS = os.environ.get('ES_CA_CERTS', '/usr/share/elasticsearch/config/certs/ca/ca.crt')
ES_HTTP_AUTH = (
    os.environ.get('ES_USER', 'elastic'), 
    os.environ.get('ES_PASSWORD', 'laylabakaa')
)

class ESIngester:
    def __init__(self, hosts=ES_HOSTS, verify_certs=True, ca_certs=ES_CA_CERTS, http_auth=ES_HTTP_AUTH):
        """Initialize the ElasticSearch Ingester with relevant information (including the BERT model)"""
        self.hosts = hosts
        self.verify_certs = verify_certs
        self.ca_certs = ca_certs
        self.http_auth = http_auth
        self._es = Elasticsearch(
                                hosts=self.hosts,
                                verify_certs=self.verify_certs,
                                ca_certs=self.ca_certs,
                                http_auth=self.http_auth
                            )
        self.pdfs = []
        self.Rs = []
        self.docs = []
        self.model = BERT()

    def setup_ingest(self, directory_path='./Layla_Notes'):
        """Iterate thru folder to sort files by type (PDFs, R-scripts, Documents)"""
        # Lists to store file paths
        paths_pdf = []
        paths_R = []
        paths_docx = []

        # Iterate through the files in the directory
        for filename in os.listdir(directory_path):
            if filename.endswith(".pdf"):
                paths_pdf.append(os.path.join(directory_path, filename))
            elif filename.endswith(".R"):
                paths_R.append(os.path.join(directory_path, filename))
            elif filename.endswith(".docx"):
                paths_docx.append(os.path.join(directory_path, filename))

        self.pdfs = paths_pdf
        self.Rs = paths_R
        self.docs = paths_docx
        
    def ingest_pdf(self, pdf_path, raw_index, vector_index):
        """Ingest PDF into ElasticSearch raw text and embeddings indices"""
        try:
            print(f'Opening {pdf_path}')
            doc = fitz.open(pdf_path)
            print('doc opened successfully')
        except Exception as e:
            print(f"Error opening PDF {pdf_path}: {e}")
            return

        for page_num, page in enumerate(doc):
            try:
                text = page.get_text()
                print(f'Page {page_num} read successfully')
                
                # Generate a unique ID for each page
                doc_id = f"{str(uuid.uuid4())}_page_{page_num}"
                
                # Generate embeddings for the page text
                embeddings = self.model.generate_embeddings(text)
                
                # Index the page text and embeddings into Elasticsearch
                self._es.index(index=raw_index, body={"content": text, "page": page_num, "pdf_path": pdf_path}, id=doc_id)
                self._es.index(index=vector_index, body={"embeddings": embeddings}, id=doc_id)
            except Exception as e:
                print(f"Error processing page {page_num} in {pdf_path}: {e}")
                continue  # Skip to the next page if there's an error

    def ingest_r_script(self, r_path, raw_index, vector_index):
        """Ingest R Script into ElasticSearch raw text and embeddings indices"""
        doc_id = str(uuid.uuid4())
        with open(r_path, 'r') as file:
            content = file.read()

        embeddings = self.model.generate_embeddings(content)

        self._es.index(index=raw_index, body={"content": content}, id=doc_id)
        self._es.index(index=vector_index, body={"embeddings": embeddings.tolist()}, id=doc_id)

    def ingest_docx(self, docx_path, raw_index, vector_index):
        """Ingest Document into ElasticSearch raw text and embeddings indices"""
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"Error opening DOCX {docx_path}: {e}")
            return
        
        for para_num, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():  # Skip empty paragraphs
                continue
            
            # Generate a unique ID for each paragraph
            doc_id = f"{str(uuid.uuid4())}_para_{para_num}"
            
            # Generate embeddings for the paragraph text
            embeddings = self.model.generate_embeddings(text)
            
            # Index the paragraph text and embeddings into Elasticsearch
            self._es.index(index=raw_index, body={"content": text, "paragraph": para_num, "docx_path": docx_path}, id=doc_id)
            self._es.index(index=vector_index, body={"embeddings": embeddings}, id=doc_id)
            print(f"Paragraph {para_num} indexed successfully")

    def create_es_text_index(self, index_name):
        """Create a raw ES index"""
        index_body = {
            "mappings": {
                "properties": {
                    "content": {"type": "text"}
                }
            }
        }
        self.create_index_if_not_exists(index_name, index_body)

    def create_es_vector_index(self, index_name, dimensions=384):
        """Create a vector ES index"""
        index_body = {
            "mappings": {
                "properties": {
                    "embeddings": {
                        "type": "dense_vector",
                        "dims": dimensions
                    }
                }
            }
        }
        self.create_index_if_not_exists(index_name, index_body)

    def create_index_if_not_exists(self, index_name, index_body):
        """Creates an ES index if it does not already exist"""
        if not self._es.indices.exists(index=index_name):
            self._es.indices.create(index=index_name, body=index_body)
            print(f"Index {index_name} created.")
        else:
            print(f"Index {index_name} already exists.")


    def ingest_docx_for_summary(self, docx_path, raw_index, vector_index):
        """Ingests DOCs for summarizer model """
        #Read document one at a time
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"Error opening DOCX {docx_path}: {e}")
            return
            
        full_text = []
        current_lecture = []
        for para in doc.paragraphs:
            #First, we should split the document into lectures
            if para.text.startswith(tuple([f"{num}." for num in range(1, 100)])):  # New lecture detection
                if current_lecture:  # Save previous lecture if exists
                    full_text.append(" ".join(current_lecture))
                    current_lecture = []  # Reset for next lecture
            #Second, we should iterate thru each lecture and remove speaker info and timestamps using regular expression
            cleaned_text = re.sub(r"\d{1,2}:\d{2}", "", para.text)  # Removes timestamps
            cleaned_text = re.sub(r"^[A-Za-z0-9\s]+:", "", cleaned_text)  # Removes speaker names
            #Third, we should remove new lines to have fluidity in text
            current_lecture.append(cleaned_text.strip())
        
        if current_lecture:  # Don't forget the last lecture
            full_text.append(" ".join(current_lecture))
        
        # Process each lecture for summarization and indexing here
        for num, lecture in enumerate(full_text):
            #Index the lecture with embeddings
            # Generate a unique ID for each paragraph
            doc_id = f"{str(uuid.uuid4())}_lec_{num}"
            
            # Generate embeddings for the paragraph text
            embeddings = self.model.generate_embeddings(lecture)
            print(lecture)
            
            # Index the paragraph text and embeddings into Elasticsearch
            self._es.index(index=raw_index, body={"content": lecture, "lec_num": num, "docx_path": docx_path}, id=doc_id)
            self._es.index(index=vector_index, body={"embeddings": embeddings}, id=doc_id)

    def natural_language_from_template(self, raw_data_path, template, raw_index, vector_index):
        """Ingest SQL data from a natural language template"""
        with open(raw_data_path) as f:
            data = json.load(f)
        
        placeholder_names = re.findall(r'\{([^}]+)\}', template)
        # Iterate through each item in the JSON data
        for item in data:
            # Create a dictionary containing only the keys present in the template
            formatted_item = {key: item[key] for key in placeholder_names if key in item}
            
            # Apply the template to the filtered item
            formatted_string = template.format_map(formatted_item)

            doc_id = f"{str(uuid.uuid4())}_app"
            embeddings = self.model.generate_embeddings(formatted_string)

            self._es.index(index=raw_index, body={"content": formatted_string}, id=doc_id)
            self._es.index(index=vector_index, body={"embeddings": embeddings}, id=doc_id)