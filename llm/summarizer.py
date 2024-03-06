from transformers import pipeline
from docx import Document
import re
import os
import json
import nltk
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

class Summarizer:
    def __init__(self):
        # Initialize the pipeline with the summarization model
        self.model = pipeline("summarization", model="facebook/bart-large-cnn")
        self.docs = []

    def setup_ingest(self, directory_path='./Layla_Notes'):
        # Lists to store file paths
        paths_docx = []

        # Iterate through the files in the directory
        for filename in os.listdir(directory_path):
            if filename.endswith(".docx"):
                paths_docx.append(os.path.join(directory_path, filename))
        self.docs = paths_docx

    def chunk_text_by_sentences(self, text, chunk_size=1024):
        sentences = sent_tokenize(text)
        current_chunk = ""
        chunks = []

        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 1 > chunk_size:
                chunks.append(current_chunk)
                current_chunk = sentence
            else:
                current_chunk += " " + sentence
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks

    def summarize(self, text):
        chunks = self.chunk_text_by_sentences(text)
        summarized_text = []
        
        for chunk in chunks:
            summary = self.model(chunk, truncation=True)
            summarized_text.append(summary[0]['summary_text'])
        
        return " ".join(summarized_text)

    def prepare_and_summarize_docx(self, output_folder):
        """Prepares DOCX for summarizer model and saves the summarized text"""
        for docx_path in self.docs:
            try:
                doc = Document(docx_path)
            except Exception as e:
                print(f"Error opening DOCX {docx_path}: {e}")
                return

            full_text = []
            current_lecture = []
            for para in doc.paragraphs:
                cleaned_text = re.sub(r"\d{1,2}:\d{2}(:\d{2})?", "", para.text)  # Removes timestamps, including those with hours
                cleaned_text = re.sub(r"^[A-Za-z0-9\s]+:", "", cleaned_text)  # Removes speaker names
                current_lecture.append(cleaned_text.strip())
            
            if current_lecture:  # Don't forget the last paragraph
                full_text.append(" ".join(current_lecture))

            full_text_str = " ".join(full_text)

            summarized_doc = Document()

            summarized_lecture = self.summarize(full_text_str)
            summarized_doc.add_paragraph(summarized_lecture)
            summarized_doc.add_page_break()

            # Save to DOCX
            try:
                output_path = os.path.join(output_folder, docx_path[14:-5] + '_SUMMARY.docx')
                print(output_path)
                summarized_doc.save(output_path)
                print(f"Summarized document saved to {output_path}")
            except Exception as e:
                print(f"Failed to save summarized document: {e}")

            # Save to JSON
            try:
                json_output_path = os.path.join(output_folder, docx_path[14:-5] + '_SUMMARY.json')
                with open(json_output_path, 'w', encoding='utf-8') as f:
                    json.dump({"summaries": summarized_lecture}, f, ensure_ascii=False, indent=4)
                print(f"Summarized text saved to {json_output_path}")
            except Exception as e:
                print(f"Failed to save summarized text to JSON: {e}")


if __name__ == "__main__":
    summarizer = Summarizer()
    summarizer.setup_ingest()
    summarizer.prepare_and_summarize_docx('./Layla_Notes_Summary')